"""
Resume Exporter - Export resumes to PDF, DOCX, and Markdown formats
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_markdown(resume_text: str, output_path: str) -> str:
    """
    Export resume to Markdown format.
    
    Args:
        resume_text: Plain text resume content
        output_path: Path to save the markdown file
        
    Returns:
        Path to the saved file
    """
    # Convert plain text to markdown with proper formatting
    lines = resume_text.split('\n')
    markdown_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            markdown_lines.append('')
            continue
            
        # Convert section headers (all caps lines) to markdown headers
        if line.isupper() and len(line) > 3:
            markdown_lines.append(f'\n## {line.title()}\n')
        # Convert bullet points
        elif line.startswith('-') or line.startswith('•'):
            clean_line = line.lstrip('-•').strip()
            markdown_lines.append(f'- {clean_line}')
        else:
            markdown_lines.append(line)
    
    markdown_content = '\n'.join(markdown_lines)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    return output_path


def export_to_docx(resume_text: str, output_path: str) -> str:
    """
    Export resume to DOCX format with professional styling.
    
    Args:
        resume_text: Plain text resume content
        output_path: Path to save the DOCX file
        
    Returns:
        Path to the saved file
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    lines = resume_text.split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            # Add empty paragraph for spacing
            doc.add_paragraph()
            continue
        
        # Check if it's a section header (all caps)
        if line.isupper() and len(line) > 3:
            # Add section header
            heading = doc.add_heading(line.title(), level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Add a line under the heading
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
        # Check if it's a bullet point
        elif line.startswith('-') or line.startswith('•'):
            clean_line = line.lstrip('-•').strip()
            para = doc.add_paragraph(clean_line, style='List Bullet')
        # Check if it's the name (first non-empty line, typically larger)
        elif i < 3 and not any(c in line for c in ['@', '|', '-', '•']):
            heading = doc.add_heading(line, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para = doc.add_paragraph(line)
    
    doc.save(output_path)
    return output_path


from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors

def export_to_pdf(resume_text: str, output_path: str) -> str:
    """
    Export resume to professional PDF format matching the reference template.
    
    Reference Format:
    - Name: Bold, centered, dark
    - Contact: Centered, gray, pipe-separated
    - Section Headers: Bold, uppercase, with underline
    - Job Entry: "Title, Company | Dates" in bold blue
    - Bullets: Clean bullet points, no asterisks
    - Summary: Normal body text (NOT bold)
    """
    import re
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    
    # =========================================================================
    # STEP 1: PREPROCESS - Clean raw AI output and remove ALL asterisks
    # =========================================================================
    def preprocess_resume(text):
        """Clean up raw AI output markers and remove asterisks."""
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # Skip empty lines at the very start
            if not cleaned_lines and not stripped:
                continue
            
            # Skip AI/humanizer output markers
            skip_patterns = [
                r'^Section:\s*',
                r'^SECTION:\s*',
                r'^Rewritten Content:\s*$',
                r'^REWRITTEN CONTENT:\s*$',
                r'^---+$',
                r'^Original Content:\s*$',
                r'^AI Probability:\s*\d+%',
                r'^Issues:\s*',
            ]
            
            should_skip = False
            for pattern in skip_patterns:
                if re.match(pattern, stripped, re.IGNORECASE):
                    should_skip = True
                    break
            
            if should_skip:
                continue
            
            # Remove ALL asterisks (they cause formatting issues)
            cleaned = stripped.replace('**', '').replace('*', '')
            
            # Convert markdown bullet to proper bullet
            if cleaned.startswith('- '):
                cleaned = '• ' + cleaned[2:]
            
            cleaned_lines.append(cleaned)
        
        return '\n'.join(cleaned_lines)
    
    # =========================================================================
    # STEP 2: DETECTION HELPERS
    # =========================================================================
    
    # Resume section headers
    SECTION_HEADERS = {
        'PROFESSIONAL SUMMARY', 'SUMMARY', 'PROFILE', 'OBJECTIVE', 'ABOUT',
        'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT', 'PROFESSIONAL EXPERIENCE',
        'EDUCATION', 'ACADEMIC BACKGROUND',
        'SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES', 'KEY SKILLS',
        'PROJECTS', 'KEY PROJECTS', 'PERSONAL PROJECTS',
        'CERTIFICATIONS', 'CERTIFICATES', 'LICENSES',
        'ACHIEVEMENTS', 'AWARDS', 'HONORS',
        'PUBLICATIONS', 'LANGUAGES', 'INTERESTS'
    }
    
    def is_section_header(line):
        """Detect if a line is a section header."""
        clean = line.strip().upper()
        return clean in SECTION_HEADERS or (
            line.isupper() and 3 < len(clean) < 35 and '|' not in line
        )
    
    def is_job_entry(line):
        """Detect if a line is a job/company entry."""
        # Pattern: "Title, Company | Date - Date" or "Company | Title | Date"
        has_pipe = '|' in line
        has_date = re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d{2}|Present)', line, re.IGNORECASE)
        return has_pipe and has_date
    
    def is_project_entry(line):
        """Detect if a line is a project name (not a bullet, not a header)."""
        # Project entries are typically short, no bullet, might have a colon
        if line.startswith('•') or line.startswith('-'):
            return False
        if is_section_header(line):
            return False
        if is_job_entry(line):
            return False
        # Short lines that end with colon or are project-like
        if ':' in line and len(line) < 60:
            return True
        return False
    
    def is_contact_info(line):
        """Detect if a line is contact information."""
        indicators = ['@', 'linkedin.com', 'github.com', '+91', '+1']
        line_lower = line.lower()
        has_indicator = any(ind in line_lower for ind in indicators)
        has_pipes = line.count('|') >= 2
        return has_indicator or (has_pipes and not re.search(r'20\d{2}', line))
    
    def is_bullet_point(line):
        """Detect if a line is a bullet point."""
        return line.startswith('•') or line.startswith('- ') or line.startswith('– ')
    
    # =========================================================================
    # STEP 3: CLEAN AND PROCESS TEXT
    # =========================================================================
    resume_text = preprocess_resume(resume_text)
    
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.6*inch,
        leftMargin=0.6*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = getSampleStyleSheet()
    
    # =========================================================================
    # STEP 4: DEFINE PROFESSIONAL STYLES (Matching Reference)
    # =========================================================================
    
    # Name style - bold, dark, left-aligned (like reference)
    name_style = ParagraphStyle(
        'ResumeName',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        alignment=TA_LEFT,
        spaceAfter=4,
        textColor=colors.HexColor('#1a1a2e')
    )
    
    # Contact info style - smaller, gray
    contact_style = ParagraphStyle(
        'ContactInfo',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        alignment=TA_LEFT,
        spaceAfter=12,
        textColor=colors.HexColor('#444444')
    )
    
    # Section header style - bold, uppercase
    section_header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11,
        spaceBefore=12,
        spaceAfter=4,
        textColor=colors.HexColor('#1a1a2e'),
    )
    
    # Job entry style - bold, dark blue (like reference)
    job_entry_style = ParagraphStyle(
        'JobEntry',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        spaceBefore=8,
        spaceAfter=2,
        textColor=colors.HexColor('#2c5282')  # Professional blue
    )
    
    # Project entry style - bold, dark (no asterisks)
    project_style = ParagraphStyle(
        'ProjectEntry',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        spaceBefore=6,
        spaceAfter=2,
        textColor=colors.HexColor('#1a1a2e')
    )
    
    # Bullet point style - proper indentation
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leftIndent=12,
        firstLineIndent=-12,
        spaceAfter=2,
        textColor=colors.HexColor('#333333'),
        alignment=TA_JUSTIFY
    )
    
    # Normal body text - for summary and other text
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        spaceAfter=4,
        textColor=colors.HexColor('#333333'),
        alignment=TA_JUSTIFY
    )
    
    # =========================================================================
    # STEP 5: BUILD THE PDF
    # =========================================================================
    story = []
    lines = resume_text.split('\n')
    
    name_found = False
    in_summary = False
    in_projects = False
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # Skip empty lines but add small spacing
        if not line:
            story.append(Spacer(1, 3))
            in_summary = False
            continue
        
        # 1. Name (first substantial line)
        if not name_found and i < 3 and not is_contact_info(line) and not is_section_header(line):
            if len(line) > 3 and not is_bullet_point(line):
                story.append(Paragraph(line, name_style))
                name_found = True
                continue
        
        # 2. Contact info
        if is_contact_info(line) and i < 6:
            story.append(Paragraph(line, contact_style))
            continue
        
        # 3. Section headers
        if is_section_header(line):
            clean_header = line.upper()
            story.append(Paragraph(clean_header, section_header_style))
            story.append(HRFlowable(
                width="100%",
                thickness=0.5,
                lineCap='round',
                color=colors.HexColor('#1a1a2e'),
                spaceBefore=1,
                spaceAfter=4
            ))
            # Track if we're in summary or projects section
            in_summary = 'SUMMARY' in clean_header
            in_projects = 'PROJECT' in clean_header
            continue
        
        # 4. Summary section - use body style (NOT bold)
        if in_summary:
            story.append(Paragraph(line, body_style))
            in_summary = False  # Only first paragraph after header
            continue
        
        # 5. Job/Company entries
        if is_job_entry(line):
            story.append(Paragraph(line, job_entry_style))
            continue
        
        # 6. Project entries (in projects section)
        if in_projects and is_project_entry(line):
            story.append(Paragraph(line, project_style))
            continue
        
        # 7. Bullet points
        if is_bullet_point(line):
            # Clean the bullet marker and add our own
            clean_bullet = line.lstrip('•-– ').strip()
            story.append(Paragraph(f"• {clean_bullet}", bullet_style))
            continue
        
        # 8. Default: body text
        story.append(Paragraph(line, body_style))
    
    # Build the PDF
    doc.build(story)
    return output_path


def export_resume(resume_text: str, job_id: str, output_dir: str) -> dict:
    """
    Export resume to all three formats (Markdown, DOCX, PDF).
    
    Args:
        resume_text: Plain text resume content
        job_id: Unique job identifier
        output_dir: Directory to save files
        
    Returns:
        Dict with paths to all exported files
    """
    os.makedirs(output_dir, exist_ok=True)
    
    paths = {}
    
    # Export Markdown
    md_path = os.path.join(output_dir, f'resume_{job_id}.md')
    paths['markdown'] = export_to_markdown(resume_text, md_path)
    
    # Export DOCX
    docx_path = os.path.join(output_dir, f'resume_{job_id}.docx')
    paths['docx'] = export_to_docx(resume_text, docx_path)
    
    # Export PDF
    pdf_path = os.path.join(output_dir, f'resume_{job_id}.pdf')
    paths['pdf'] = export_to_pdf(resume_text, pdf_path)
    
    return paths
